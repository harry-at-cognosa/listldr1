"""
Proof of Concept: Section Swapping in Word Documents

This demonstrates replacing a section in a document with content from another file.
"""

from docx import Document
from docx.oxml.ns import qn
from copy import deepcopy
import re
import os


def find_section_element_index(doc, section_num: int) -> int | None:
    """
    Find the element index where a section header starts.
    Looks for patterns like '4 – Product Pump' in tables and paragraphs.
    """
    section_pattern = re.compile(rf'^{section_num}\s*[–\-]\s*')
    body = doc.element.body

    for idx, child in enumerate(body):
        tag = child.tag.split('}')[-1]

        if tag == 'p':
            text = ''.join(t.text or '' for t in child.findall('.//' + qn('w:t')))
            if section_pattern.match(text.strip()):
                return idx

        elif tag == 'tbl':
            # Check first cell
            first_cell_texts = child.findall('.//' + qn('w:tc'))[0].findall('.//' + qn('w:t'))
            text = ''.join(t.text or '' for t in first_cell_texts)
            if section_pattern.match(text.strip()):
                return idx

    return None


def replace_section(doc_path: str, section_num: int, replacement_path: str, output_path: str):
    """
    Replace a section in a document with content from a replacement file.

    Args:
        doc_path: Path to the original document
        section_num: The section number to replace (e.g., 4 for 'Section 4 - Product Pump')
        replacement_path: Path to the docx file containing replacement content
        output_path: Where to save the modified document
    """
    doc = Document(doc_path)
    replacement = Document(replacement_path)

    # Find where section N starts
    start_idx = find_section_element_index(doc, section_num)
    if start_idx is None:
        raise ValueError(f"Section {section_num} not found in document")

    # Find where section N+1 starts (or end of document)
    end_idx = find_section_element_index(doc, section_num + 1)
    if end_idx is None:
        # If no next section, look for higher numbers
        for next_num in range(section_num + 2, 20):
            end_idx = find_section_element_index(doc, next_num)
            if end_idx is not None:
                break

    if end_idx is None:
        # No next section found - go to end but leave footer
        # (This is simplified - in production you'd detect footer more carefully)
        end_idx = len(list(doc.element.body)) - 5  # Leave last 5 elements as footer

    print(f"Section {section_num} spans elements {start_idx} to {end_idx - 1}")
    print(f"Removing {end_idx - start_idx} elements")

    # Get references to elements to remove
    body = doc.element.body
    elements_to_remove = list(body)[start_idx:end_idx]

    # Remove old section elements
    for elem in elements_to_remove:
        body.remove(elem)

    print(f"Removed old section {section_num}")

    # Insert replacement content at the same position
    # Get insertion point (element now at start_idx, or end if we removed everything)
    body_children = list(body)
    if start_idx < len(body_children):
        insert_before = body_children[start_idx]
    else:
        insert_before = None

    # Copy elements from replacement document
    replacement_body = replacement.element.body
    elements_to_insert = list(replacement_body)

    print(f"Inserting {len(elements_to_insert)} elements from replacement")

    for elem in elements_to_insert:
        new_elem = deepcopy(elem)
        if insert_before is not None:
            insert_before.addprevious(new_elem)
        else:
            body.append(new_elem)

    # Save the result
    doc.save(output_path)
    print(f"Saved modified document to: {output_path}")


def extract_section(doc_path: str, section_num: int, output_path: str):
    """
    Extract a single section from a document to a new file.
    Useful for creating section files that can be used as replacements.
    """
    doc = Document(doc_path)

    start_idx = find_section_element_index(doc, section_num)
    if start_idx is None:
        raise ValueError(f"Section {section_num} not found")

    # Find end
    end_idx = None
    for next_num in range(section_num + 1, 20):
        end_idx = find_section_element_index(doc, next_num)
        if end_idx is not None:
            break

    if end_idx is None:
        end_idx = len(list(doc.element.body)) - 5

    print(f"Extracting section {section_num}: elements {start_idx} to {end_idx - 1}")

    # Create new document with just this section
    new_doc = Document()
    new_body = new_doc.element.body

    # Clear default empty paragraph
    for child in list(new_body):
        new_body.remove(child)

    # Copy section elements
    body = doc.element.body
    for idx, child in enumerate(body):
        if start_idx <= idx < end_idx:
            new_body.append(deepcopy(child))

    new_doc.save(output_path)
    print(f"Extracted section {section_num} to: {output_path}")


# =============================================================================
# Demo
# =============================================================================

if __name__ == "__main__":
    source = '/Users/harry/1_listldr_work/UBM_20_mils/UBM 20 FCFC_US_FOB Allendale_11 2025.docx'

    # Make sure output directory exists
    os.makedirs('outputs', exist_ok=True)

    print("=" * 60)
    print("DEMO: Extract Section 4 (Product Pump)")
    print("=" * 60)
    extract_section(source, section_num=4, output_path='outputs/section_4_product_pump.docx')

    print()
    print("=" * 60)
    print("DEMO: Extract Section 6 (Grinding Media)")
    print("=" * 60)
    extract_section(source, section_num=6, output_path='outputs/section_6_grinding_media.docx')

    print()
    print("Now you can:")
    print("1. Edit outputs/section_4_product_pump.docx in Word")
    print("2. Use replace_section() to swap it into a document")
    print()
    print("Example replacement (uncomment to run):")
    print("  replace_section(source, 4, 'outputs/section_4_product_pump.docx', 'outputs/modified.docx')")
