"""
Proof of Concept: Document Assembly with docxcompose

This demonstrates:
1. Assembling a document from separate section files
2. Parsing an existing document to extract sections by numbered headings

Requirements:
    pip install python-docx docxcompose
"""

from docx import Document
from docxcompose.composer import Composer
import re
import os


def assemble_document(cover_page_path: str, section_paths: list[str], output_path: str):
    """
    Assemble a complete document from a cover page and section files.

    The cover page serves as the 'vessel' - it contains all the trade dress,
    watermarks, headers/footers, etc. Section documents are appended to it.

    Args:
        cover_page_path: Path to the cover page docx (contains formatting/branding)
        section_paths: List of paths to section docx files, in order
        output_path: Where to save the assembled document
    """
    # Start with the cover page as the base - this preserves its formatting
    master = Document(cover_page_path)
    composer = Composer(master)

    # Append each section in order
    for section_path in section_paths:
        section_doc = Document(section_path)
        composer.append(section_doc)
        print(f"  Added: {os.path.basename(section_path)}")

    # Save the assembled document
    composer.save(output_path)
    print(f"\nAssembled document saved to: {output_path}")


def parse_sections_from_template(template_path: str, output_dir: str):
    """
    Parse an existing template document into separate section files.

    Looks for numbered section headings like:
        "1 – Principal Characteristics"
        "2 – General Technical Data"

    Each section (from one heading to the next) is saved as a separate docx.

    Args:
        template_path: Path to the source template docx
        output_dir: Directory to save extracted section files
    """
    doc = Document(template_path)

    # Pattern to match section headings: number, dash/hyphen variants, title
    # Matches: "1 – Title", "1 - Title", "10 – Title", etc.
    section_pattern = re.compile(r'^(\d+)\s*[–\-]\s*(.+)$')

    os.makedirs(output_dir, exist_ok=True)

    # Track sections as we find them
    sections = []  # List of (section_num, section_title, start_index)

    # First pass: identify where each section starts
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        match = section_pattern.match(text)
        if match:
            section_num = match.group(1)
            section_title = match.group(2).strip()
            sections.append((section_num, section_title, i))
            print(f"Found section {section_num}: {section_title} (paragraph {i})")

    if not sections:
        print("No numbered sections found in document.")
        return

    # Find cover page content (everything before first section)
    first_section_start = sections[0][2]
    if first_section_start > 0:
        print(f"\nCover page: paragraphs 0-{first_section_start - 1}")
        _save_paragraph_range(doc, 0, first_section_start,
                              os.path.join(output_dir, "00_cover_page.docx"))

    # Extract each section
    for i, (section_num, section_title, start_idx) in enumerate(sections):
        # End index is start of next section, or end of document
        if i + 1 < len(sections):
            end_idx = sections[i + 1][2]
        else:
            end_idx = len(doc.paragraphs)

        # Create safe filename
        safe_title = re.sub(r'[^\w\s-]', '', section_title).strip()
        safe_title = re.sub(r'\s+', '_', safe_title)
        filename = f"{section_num.zfill(2)}_{safe_title}.docx"
        filepath = os.path.join(output_dir, filename)

        _save_paragraph_range(doc, start_idx, end_idx, filepath)
        print(f"Saved: {filename} ({end_idx - start_idx} paragraphs)")


def _save_paragraph_range(source_doc: Document, start_idx: int, end_idx: int, output_path: str):
    """
    Save a range of paragraphs from source document to a new document.

    Note: This is a simplified extraction. It copies paragraph text and
    attempts to preserve basic formatting, but complex elements (tables,
    images, etc.) require additional handling.
    """
    new_doc = Document()

    for i in range(start_idx, end_idx):
        source_para = source_doc.paragraphs[i]

        # Create new paragraph with same style
        new_para = new_doc.add_paragraph()
        new_para.style = source_para.style

        # Copy runs (text segments with formatting)
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            # Copy run formatting
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.name:
                new_run.font.name = run.font.name

    new_doc.save(output_path)


def list_document_structure(template_path: str):
    """
    Utility: Print the structure of a document to understand its layout.
    Shows paragraphs with their styles and first 60 chars of text.
    """
    doc = Document(template_path)

    print(f"\nDocument structure: {template_path}")
    print("=" * 70)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()[:60]
        style = para.style.name if para.style else "None"
        if text:  # Only show non-empty paragraphs
            print(f"{i:3d} [{style:20s}] {text}...")

    # Also count tables
    print(f"\nTables in document: {len(doc.tables)}")


# =============================================================================
# Example usage
# =============================================================================

if __name__ == "__main__":
    print("DocxCompose Proof of Concept")
    print("=" * 40)

    # Example 1: Assemble a document from parts
    # (Uncomment and adjust paths when you have actual section files)
    """
    assemble_document(
        cover_page_path="sections/00_cover_page.docx",
        section_paths=[
            "sections/01_Principal_Characteristics.docx",
            "sections/02_General_Technical_Data.docx",
            "sections/03_Product_Pump.docx",
            "sections/04_Options.docx",
        ],
        output_path="outputs/assembled_template.docx"
    )
    """

    # Example 2: Parse an existing template into sections
    # (Uncomment and adjust path when you have a template to parse)
    """
    parse_sections_from_template(
        template_path="inputs/existing_template.docx",
        output_dir="outputs/extracted_sections"
    )
    """

    # Example 3: Inspect document structure
    # (Uncomment and adjust path to examine any docx)
    """
    list_document_structure("inputs/existing_template.docx")
    """

    print("\nTo test, uncomment the examples above and provide real docx files.")
    print("Place input files in the 'inputs/' folder.")
